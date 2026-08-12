from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Propuesta_Generador_Imagenes_OpenAI.docx"
BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(23)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor.from_string(BLUE)
for style_name in ("Heading 1", "Heading 2"):
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor.from_string(BLUE)
styles["Heading 1"].font.size = Pt(15)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Propuesta: Generador de Imágenes por Lote")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Automatización de edición de imágenes mediante OpenAI API")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(89, 89, 89)
date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date.add_run("Presentación ejecutiva · 11 de agosto de 2026").font.size = Pt(9)

doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    "El proyecto permite procesar automáticamente carpetas completas de imágenes de productos. "
    "Cada archivo se envía de forma individual al modelo GPT Image 2 con una instrucción común "
    "—por ejemplo, limpiar el fondo, uniformar iluminación o preparar una fotografía de catálogo— "
    "y el resultado se guarda en formato PNG."
)

doc.add_heading("Beneficio para el negocio", level=1)
for item in (
    "Reduce el trabajo manual repetitivo de edición.",
    "Mantiene instrucciones y presentación consistentes en todo el catálogo.",
    "Permite aprobar una vista previa antes de procesar el lote completo.",
    "Muestra una estimación del gasto antes de autorizar cada ejecución.",
):
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Flujo de trabajo", level=1)
for item in (
    "Seleccionar la carpeta con imágenes JPG, PNG o WEBP.",
    "Definir la instrucción de edición y el tamaño de salida.",
    "Generar y revisar una imagen de prueba.",
    "Aprobar el lote y guardar los resultados con el sufijo “_generado”.",
):
    doc.add_paragraph(item, style="List Number")

doc.add_heading("Tarifa oficial de GPT Image 2", level=1)
doc.add_paragraph(
    "Coste estimado de la imagen de salida, en dólares estadounidenses (USD), según calidad y tamaño:"
)
table = doc.add_table(rows=4, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["Calidad", "1024 × 1024", "1024 × 1536", "1536 × 1024"]
rows = [
    ["Baja", "$0.006", "$0.005", "$0.005"],
    ["Media", "$0.053", "$0.041", "$0.041"],
    ["Alta", "$0.211", "$0.165", "$0.165"],
]
for i, value in enumerate(headers):
    shade(table.cell(0, i), BLUE)
    set_cell_text(table.cell(0, i), value, bold=True, color="FFFFFF")
for r_idx, values in enumerate(rows, start=1):
    for c_idx, value in enumerate(values):
        if r_idx % 2:
            shade(table.cell(r_idx, c_idx), LIGHT_BLUE)
        set_cell_text(table.cell(r_idx, c_idx), value, bold=(c_idx == 0))

note = doc.add_paragraph()
note.add_run("Importante: ").bold = True
note.add_run(
    "al editar fotografías existentes también se cobran los tokens de la imagen de entrada y del texto. "
    "Por eso el coste final puede ser ligeramente superior a la tabla. La aplicación utiliza una "
    "estimación conservadora de $0.06 por imagen para planificación en calidad media."
)

doc.add_heading("Presupuesto orientativo del proyecto", level=1)
budget = doc.add_table(rows=5, cols=3)
budget.alignment = WD_TABLE_ALIGNMENT.CENTER
budget.style = "Table Grid"
budget_headers = ["Cantidad", "Estimado unitario", "Total estimado"]
budget_rows = [
    ["25 imágenes (prueba)", "$0.06", "$1.50"],
    ["100 imágenes", "$0.06", "$6.00"],
    ["500 imágenes", "$0.06", "$30.00"],
    ["1,000 imágenes", "$0.06", "$60.00"],
]
for i, value in enumerate(budget_headers):
    shade(budget.cell(0, i), BLUE)
    set_cell_text(budget.cell(0, i), value, bold=True, color="FFFFFF")
for r_idx, values in enumerate(budget_rows, start=1):
    for c_idx, value in enumerate(values):
        if r_idx % 2 == 0:
            shade(budget.cell(r_idx, c_idx), LIGHT_GRAY)
        set_cell_text(budget.cell(r_idx, c_idx), value, bold=(c_idx == 2))

doc.add_heading("Recomendación", level=1)
doc.add_paragraph(
    "Autorizar una prueba inicial de 25 imágenes, con un presupuesto operativo de USD 2 a 3 para "
    "absorber variaciones por imágenes de entrada. Después de validar calidad, tiempo ahorrado y coste "
    "real, ampliar el uso por lotes con un límite mensual definido en la cuenta de OpenAI."
)

doc.add_heading("Control y transparencia", level=1)
for item in (
    "La vista previa utiliza una llamada pagada, pero se reutiliza como primera imagen del lote si se aprueba.",
    "El redimensionamiento final a 1080 × 1080 se realiza localmente y no añade coste de API.",
    "Los errores se registran por imagen y no detienen el resto del lote.",
    "El gasto real puede consultarse en https://platform.openai.com/usage.",
):
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Fuentes", level=1)
p = doc.add_paragraph()
p.add_run("Precios oficiales y calculadora de generación de imágenes: ").bold = True
p.add_run("https://developers.openai.com/api/docs/guides/image-generation#calculating-costs")
p = doc.add_paragraph()
p.add_run("Tarifas generales de la API: ").bold = True
p.add_run("https://developers.openai.com/api/docs/pricing")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Generador de Imágenes por Lote · Propuesta de uso y presupuesto").font.size = Pt(8)

doc.save(OUTPUT)
print(OUTPUT)
